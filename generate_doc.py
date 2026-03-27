import os
import json
import time
import shutil
import glob
from datetime import datetime
from openai import OpenAI
from docx import Document

# Import existing models and sections
try:
    from src.models import SessionReport, SessionMetadata, DetailItem, SummaryItem
    from src.sections import add_header, add_summary, add_details, add_next_steps, add_footer
    from src.startup import preflight
except ImportError:
    import sys
    sys.path.append(os.getcwd())
    from src.models import SessionReport, SessionMetadata, DetailItem, SummaryItem
    from src.sections import add_header, add_summary, add_details, add_next_steps, add_footer
    from src.startup import preflight

# ── Configuration ────────────────────────────────────────────────────────
LM_STUDIO_URL = "http://localhost:1234/v1"

# ── Directories ──────────────────────────────────────────────────────────
INPUT_DIR = "input_docs"
PROCESSED_DIR = "processed_docs"
OUTPUT_DIR = "output_docs"

# ── System Prompt ────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are a clinical scribe. Extract ALL information from the session notes into exactly this JSON structure. Output ONLY valid JSON, nothing else.

{
  "metadata": {
    "patient_name": "LASTNAME, Firstname",
    "dob": "date of birth or N/A",
    "session_date": "full date",
    "session_time": "time",
    "session_type": "type of session"
  },
  "summary_intro": "Opening summary paragraph that gives a high-level overview of the patient.",
  "summary_sections": [
    {"heading": "Sub-section heading", "content": "Full paragraph for this summary sub-section"}
  ],
  "details": [
    {"heading": "Sub-section heading", "content": "Full paragraph for this section"}
  ],
  "next_steps": ["Each suggested next step as a separate string"]
}

RULES — follow every one or your output is invalid:
1. Every object in "summary_sections" and "details" MUST have BOTH "heading" AND "content" keys.
2. The "content" value must be a non-empty string with complete sentences extracted from the notes.
3. STRICT EXTRACT ONLY: Do NOT add any new information, interpretations, or assumptions. Extract ONLY what is explicitly stated in the text.
4. STRICT NO OMISSION: Do NOT remove, omit, or summarize away any information, sections, or details from the source document.
5. Do NOT produce an object that only has a "heading" key with no "content" key.
6. Output ONLY the JSON object — no markdown fences, no comments, no extra text."""


# ── Helpers ──────────────────────────────────────────────────────────────

def fmt_time(seconds):
    """Format seconds into Xm Ys string."""
    mins = int(seconds) // 60
    secs = seconds % 60
    if mins > 0:
        return f"{mins}m {secs:.2f}s"
    return f"{secs:.2f}s"


def log(message: str, style: str = "info"):
    """Structured logging with timestamps."""
    timestamp = datetime.now().strftime("%H:%M:%S")
    if style == "header":
        print(f"\n[{timestamp}] ════ {message} ════")
    elif style == "success":
        print(f"[{timestamp}] ✓ {message}")
    elif style == "error":
        print(f"[{timestamp}] ✗ {message}")
    elif style == "info":
        print(f"[{timestamp}]   {message}")
    elif style == "dim":
        print(f"[{timestamp}]     {message}")


def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a Word document."""
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        log(f"Error reading {os.path.basename(file_path)}: {e}", "error")
        return ""


def process_file(file_path: str, client: OpenAI, model_id: str) -> SessionReport | None:
    """Send a single document to the LLM and return a validated SessionReport."""
    text = extract_text_from_docx(file_path)
    if not text:
        return None

    log(f"Loaded {len(text)} characters", "info")
    log(f"Sending to {model_id}...", "info")

    llm_start = time.time()
    try:
        response = client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": f"Extract the following session notes:\n\n{text}"}
            ],
            temperature=0.0
        )
    except Exception as e:
        log(f"LLM error: {e}", "error")
        return None

    llm_time = time.time() - llm_start
    log(f"LLM response received in {fmt_time(llm_time)}", "success")

    raw = response.choices[0].message.content.strip()

    # Strip markdown fences if model wraps in ```json
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
        raw = raw.rsplit("```", 1)[0]

    # Parse and validate
    try:
        data = json.loads(raw)
        report = SessionReport.model_validate(data)
        empty_sections = [s.heading for s in report.summary_sections + report.details if not s.content.strip()]
        if empty_sections:
            log(f"Warning: {len(empty_sections)} section(s) have empty content: {empty_sections[:3]}{'...' if len(empty_sections) > 3 else ''}", "error")
        log(f"Parsed: {len(report.summary_sections)} summary sections, "
            f"{len(report.details)} detail sections, "
            f"{len(report.next_steps)} next steps", "success")
        return report
    except json.JSONDecodeError as e:
        log(f"JSON decode error: {e}", "error")
        return None
    except Exception as e:
        log(f"Validation error: {e}", "error")
        return None


def generate_output(report: SessionReport, output_dir: str, source_filename: str):
    """Generate the Word doc and JSON from a validated report."""
    doc_start = time.time()
    out_doc = Document()
    add_header(out_doc, report.metadata)
    add_summary(out_doc, report.summary_intro, report.summary_sections)
    add_details(out_doc, report.details)
    add_next_steps(out_doc, report.next_steps)
    add_footer(out_doc)

    # Build output filename from source name
    base_name = os.path.splitext(source_filename)[0]
    docx_output = os.path.join(output_dir, f"{base_name}_Report.docx")
    json_output = os.path.join(output_dir, f"{base_name}_Report.json")

    out_doc.save(docx_output)
    doc_time = time.time() - doc_start
    log(f"Report saved in {fmt_time(doc_time)}: {os.path.basename(docx_output)}", "success")

    # Save JSON alongside
    # with open(json_output, "w") as f:
    #     json.dump(report.model_dump(), f, indent=2)
    # log(f"JSON saved: {os.path.basename(json_output)}", "dim")


# ── Main Batch Pipeline ─────────────────────────────────────────────────

def process_batch(model_id: str):
    """Batch process all .docx files in INPUT_DIR."""
    # Ensure directories exist
    os.makedirs(INPUT_DIR, exist_ok=True)
    os.makedirs(PROCESSED_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Find all docx files
    files = glob.glob(os.path.join(INPUT_DIR, "*.docx"))

    print(f"\nScanning '{INPUT_DIR}'...")
    if not files:
        print(f"No .docx files found in '{INPUT_DIR}/'. Place documents there and run again.\n")
        return

    print(f"Found {len(files)} document(s).\n")

    client = OpenAI(base_url=LM_STUDIO_URL, api_key="lm-studio")
    total_start = time.time()

    for i, file_path in enumerate(files, 1):
        filename = os.path.basename(file_path)
        log(f"Processing File {i}/{len(files)}: {filename}", "header")

        file_start = time.time()
        report = process_file(file_path, client, model_id)

        if report:
            # Generate output docs
            generate_output(report, OUTPUT_DIR, filename)

            # Move source to processed
            try:
                destination = os.path.join(PROCESSED_DIR, filename)
                if os.path.exists(destination):
                    base, ext = os.path.splitext(filename)
                    ts = time.strftime("%Y%m%d-%H%M%S")
                    destination = os.path.join(PROCESSED_DIR, f"{base}_{ts}{ext}")

                shutil.move(file_path, destination)
                log(f"Moved source → processed/", "dim")
            except Exception as e:
                log(f"Error moving file: {e}", "error")

            file_time = time.time() - file_start
            log(f"File completed in {fmt_time(file_time)}", "success")
        else:
            log(f"Failed to process {filename}", "error")

    total_time = time.time() - total_start
    print(f"\n{'='*50}")
    print(f"  Batch complete: {len(files)} file(s) in {fmt_time(total_time)}")
    print(f"  Output folder:  {os.path.abspath(OUTPUT_DIR)}")
    print(f"{'='*50}\n")


if __name__ == "__main__":
    model_id = preflight()
    process_batch(model_id)
