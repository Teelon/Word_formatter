"""Suggested Next Steps section: label + bullet list."""

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _disable_contextual_spacing(paragraph):
    """Explicitly disable contextual spacing so Word honours space before/after."""
    from lxml import etree
    pPr = paragraph._p.get_or_add_pPr()
    # Remove any existing contextualSpacing elements first
    for cs in pPr.findall(qn('w:contextualSpacing')):
        pPr.remove(cs)
    # Explicitly set to val="0" to override the style-level setting
    cs_elem = etree.SubElement(pPr, qn('w:contextualSpacing'))
    cs_elem.set(qn('w:val'), '0')


def add_next_steps(doc, next_steps: list[str]):
    """Adds the Suggested Next Steps label and items to the document body."""
    steps_label = doc.add_paragraph()
    steps_label.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    steps_label.paragraph_format.space_after = Pt(0)
    run_label = steps_label.add_run("Suggested next steps")
    run_label.bold = False
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(9)

    if not next_steps:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run("No suggested next steps were found for this meeting.")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
    else:
        for step in next_steps:
            sp = doc.add_paragraph(style='List Bullet')
            sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            sp.paragraph_format.left_indent = Inches(0.5)
            sp.paragraph_format.space_before = Pt(6)
            sp.paragraph_format.space_after = Pt(6)
            _disable_contextual_spacing(sp)
            # Clear the auto-generated run and re-add with correct font
            sp.clear()
            run = sp.add_run(step)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
