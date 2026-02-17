"""Page footer section: static clinician credentials on all pages."""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_footer(doc):
    """Adds a two-line left-aligned footer with clinician names."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    footer_para1 = footer.paragraphs[0]
    footer_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para1.paragraph_format.space_before = Pt(0)
    footer_para1.paragraph_format.space_after = Pt(0)
    f_run1 = footer_para1.add_run("G. Maynard MSc. (Sup. Prac.)")
    f_run1.font.name = 'Times New Roman'
    f_run1.font.size = Pt(9)

    footer_para2 = footer.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para2.paragraph_format.space_before = Pt(0)
    footer_para2.paragraph_format.space_after = Pt(0)
    f_run2 = footer_para2.add_run("G. Townsend, MSc. C. Psych. PsyD")
    f_run2.font.name = 'Times New Roman'
    f_run2.font.size = Pt(9)
