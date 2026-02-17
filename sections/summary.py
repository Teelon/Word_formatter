"""Summary section: label + summary text block."""

from docx.shared import Pt


def add_summary(doc, summary_text: str):
    """Adds the Summary label and text to the document body."""
    summary_label = doc.add_paragraph()
    summary_label.paragraph_format.space_after = Pt(0)
    run_label = summary_label.add_run("Summary")
    run_label.bold = False
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(9)

    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_before = Pt(0)
    run_text = summary_para.add_run(summary_text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(9)

    doc.add_paragraph()  # Spacer
