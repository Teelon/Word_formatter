"""Details section: label + indented bullet points with bold headings."""

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from models import DetailItem


def _disable_contextual_spacing(paragraph):
    """Explicitly disable contextual spacing so Word honours space before/after."""
    from lxml import etree
    pPr = paragraph._p.get_or_add_pPr()
    # Remove any existing contextualSpacing elements first
    for cs in pPr.findall(qn('w:contextualSpacing')):
        pPr.remove(cs)
    # Explicitly set to val="0" to override the style-level setting
    cs_elem = etree.SubElement(pPr, qn('w:contextualSpacing'))
    cs_elem.set(qn('w:val'), '0')


def add_details(doc, details: list[DetailItem]):
    """Adds the Details label and bullet-point items to the document body."""
    details_label = doc.add_paragraph()
    details_label.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    details_label.paragraph_format.space_after = Pt(0)
    run_label = details_label.add_run("Details")
    run_label.bold = False
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(9)

    for item in details:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        _disable_contextual_spacing(p)

        heading_run = p.add_run(f"{item.heading} ")
        heading_run.bold = True
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(9)

        content_run = p.add_run(item.content)
        content_run.font.name = 'Times New Roman'
        content_run.font.size = Pt(9)

    doc.add_paragraph()  # Spacer
