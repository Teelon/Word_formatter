"""Page header section: patient info (left) + date/time (right) + session type."""

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from models import SessionMetadata

def add_header(doc, meta: SessionMetadata):
    """Adds a two-line page header using tab stops for alignment (no table)."""
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Remove default empty paragraphs
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    # ── Line 1: patient info (left) + date/time (right via tab stop) ─────────
    line1 = header.add_paragraph()
    line1.paragraph_format.space_before = Pt(0)
    line1.paragraph_format.space_after = Pt(0)

    # Add a right-aligned tab stop at the right margin (6.5 inches)
    _add_right_tab_stop(line1, Inches(6.5))

    # Left side
    patient_info = meta.patient_name
    if meta.dob:
        patient_info += f" ({meta.dob})"
    _styled_run(line1, patient_info)

    # Tab then right-aligned date/time
    line1.add_run('\t')
    _styled_run(line1, f"{meta.session_date}; {meta.session_time}")

    # ── Line 2: session type (bold) ───────────────────────────────────────────
    line2 = header.add_paragraph()
    line2.paragraph_format.space_before = Pt(0)
    line2.paragraph_format.space_after = Pt(0)
    _styled_run(line2, meta.session_type, bold=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _styled_run(para, text: str, bold: bool = False, size: int = 9):
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return run


def _add_right_tab_stop(para, position):
    """Adds a right-aligned tab stop to a paragraph at the given position."""
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(position.pt * 20)))  # EMUs → twips
    tabs.append(tab)
    pPr.append(tabs)