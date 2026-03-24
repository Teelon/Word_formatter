"""Summary section: label + intro paragraph + dynamic sub-sections."""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_summary(doc, summary_intro: str, summary_sections: list):
    """Adds the Summary label, intro text, and sub-sections to the document body.

    Args:
        doc: The python-docx Document object.
        summary_intro: The opening summary paragraph text.
        summary_sections: List of SummaryItem objects (each with .heading and .content).
    """
    # "Summary" label
    summary_label = doc.add_paragraph()
    summary_label.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_label.paragraph_format.space_after = Pt(0)
    run_label = summary_label.add_run("Summary")
    run_label.bold = False
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(9)

    # Intro paragraph
    intro_para = doc.add_paragraph()
    intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_para.paragraph_format.space_before = Pt(0)
    intro_para.paragraph_format.space_after = Pt(4)
    run_intro = intro_para.add_run(summary_intro)
    run_intro.font.name = 'Times New Roman'
    run_intro.font.size = Pt(9)

    # Dynamic sub-sections (each with bold heading + content)
    for item in summary_sections:
        # Heading paragraph (left aligned)
        head_para = doc.add_paragraph()
        head_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        head_para.paragraph_format.space_before = Pt(2)
        head_para.paragraph_format.space_after = Pt(0)

        run_heading = head_para.add_run(item.heading)
        run_heading.bold = True
        run_heading.font.name = 'Times New Roman'
        run_heading.font.size = Pt(9)

        # Content paragraph (justified)
        content_para = doc.add_paragraph()
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content_para.paragraph_format.space_before = Pt(0)
        content_para.paragraph_format.space_after = Pt(2)

        run_content = content_para.add_run(item.content)
        run_content.font.name = 'Times New Roman'
        run_content.font.size = Pt(9)

    doc.add_paragraph()  # Spacer
